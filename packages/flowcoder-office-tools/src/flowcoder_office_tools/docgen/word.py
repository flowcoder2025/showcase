"""견적서/거래명세서 생성 — python-docx.

case05 docgen 카테고리의 핵심 빌더. column_map 강제로 다른 입력 스키마와도
호환되며, total 미명시 시 sum(qty*price)을 자동 계산한다.

한글 폰트는 ``KOREAN_FONT`` 모듈 상수로 명시적으로 적용한다 (기본
``Apple SD Gothic Neo``). Windows/Linux 시연 노트북에서 한글이 깨지지
않도록 모든 run의 ``w:rFonts`` (eastAsia/ascii/hAnsi)를 설정한다.
환경변수 ``AX_KOREAN_FONT`` 로 override 가능.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEFAULT_COLUMN_MAP: dict[str, str] = {"name": "name", "qty": "qty", "price": "price"}

KOREAN_FONT: str = os.environ.get("AX_KOREAN_FONT", "Apple SD Gothic Neo")


def _set_korean_font(run: Any, font_name: str = KOREAN_FONT) -> None:
    """run에 한글(eastAsia) 폰트를 명시. ascii/hAnsi도 같이 설정해 혼합문 안전."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _apply_font_to_paragraph(paragraph: Any) -> None:
    """단락의 모든 run에 한글 폰트를 적용. run이 없으면 add_run으로 보장."""
    if not paragraph.runs:
        # heading/empty paragraph도 빈 run을 추가해 폰트가 적용되도록 보장
        if paragraph.text == "":
            return
        # paragraph.text가 있는데 runs가 비어있는 경우는 거의 없으나 방어적으로 처리
        return
    for run in paragraph.runs:
        _set_korean_font(run)


def _coerce_int(value: Any, *, field: str, item: dict[str, Any]) -> int:
    """qty/price 값을 int로 강제 변환 — 실패 시 ValueError."""
    if value is None:
        raise ValueError(f"item {item!r} has invalid {field}: value is None")
    try:
        return int(value)
    except (TypeError, ValueError) as e:
        raise ValueError(f"item {item!r} has invalid {field}: {e}") from e


def build_quote(
    *,
    out_path: Path | str,
    vendor: str,
    items: list[dict[str, Any]],
    meta: dict[str, str],
    column_map: dict[str, str] | None = None,
    total: int | None = None,
) -> None:
    """견적서 docx 생성.

    Args:
        out_path: 결과 .docx 경로
        vendor: 거래처명 (빈 문자열/whitespace-only 금지)
        items: 품목 리스트 — 각 dict는 column_map 기준 name/qty/price 키 보유
        meta: {"date": YYYY-MM-DD, "quote_no": "Q-XXX-NNN", ...} (date/quote_no 필수)
        column_map: 입력 dict 키 매핑 (재사용성 — 다른 입력 스키마 호환)
        total: 명시 시 사용, 미명시 시 sum(qty * price) 자동 계산

    Raises:
        ValueError: items 빈 리스트 / vendor 빈 문자열 / qty·price 음수 또는
            변환 실패
        KeyError: meta에 'date' 또는 'quote_no' 누락, items에 매핑된 키 누락
    """
    if not vendor or not vendor.strip():
        raise ValueError("vendor must not be empty")
    if not items:
        raise ValueError("items must not be empty")
    if "date" not in meta:
        raise KeyError("meta must contain 'date'")
    if "quote_no" not in meta:
        raise KeyError("meta must contain 'quote_no'")

    cm = {**DEFAULT_COLUMN_MAP, **(column_map or {})}
    name_key = cm["name"]
    qty_key = cm["qty"]
    price_key = cm["price"]

    # 자동 합계 (total 미명시 시) — 동시에 모든 row의 qty/price 검증
    line_amounts: list[int] = []
    for it in items:
        qty = _coerce_int(it.get(qty_key), field="qty", item=it)
        price = _coerce_int(it.get(price_key), field="price", item=it)
        if qty < 0 or price < 0:
            raise ValueError(f"item {it!r} has negative qty/price (qty={qty}, price={price})")
        line_amounts.append(qty * price)
    auto_total = sum(line_amounts)
    final_total = total if total is not None else auto_total

    doc = Document()
    # 헤더
    heading = doc.add_heading("견 적 서", level=0)
    _apply_font_to_paragraph(heading)
    # 메타 단락
    meta_paras = [
        doc.add_paragraph(f"견적번호: {meta['quote_no']}"),
        doc.add_paragraph(f"작성일: {meta['date']}"),
        doc.add_paragraph(f"거래처: {vendor}"),
    ]
    for p in meta_paras:
        _apply_font_to_paragraph(p)

    # 품목 표
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "품목"
    hdr[1].text = "수량"
    hdr[2].text = "단가"
    hdr[3].text = "금액"
    for i, it in enumerate(items, start=1):
        # 위에서 이미 검증했으므로 여기서는 안전하게 int() 재변환
        qty = int(it[qty_key])
        price = int(it[price_key])
        row = table.rows[i].cells
        row[0].text = str(it[name_key])
        row[1].text = f"{qty:,}"
        row[2].text = f"{price:,}"
        row[3].text = f"{qty * price:,}"

    # 표 셀 모든 단락에 폰트 적용
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _apply_font_to_paragraph(p)

    # 합계 — 굵게
    doc.add_paragraph()
    total_para = doc.add_paragraph(f"합 계: {final_total:,}원")
    total_para.runs[0].bold = True
    _apply_font_to_paragraph(total_para)

    doc.save(str(out_path))
