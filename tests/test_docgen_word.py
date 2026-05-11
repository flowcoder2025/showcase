"""Tests for core.docgen.word — 견적서/거래명세서 docx 빌더."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from flowcoder_office_tools.docgen import word


def _read_full_text(path: Path) -> str:
    """docx 파일을 다시 열어 paragraph + table cell 텍스트를 합쳐 반환."""
    doc = Document(str(path))
    chunks: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def _sample_items() -> list[dict[str, object]]:
    return [
        {"name": "위젯", "qty": 10, "price": 50000},
        {"name": "기어", "qty": 5, "price": 80000},
    ]


def _sample_meta() -> dict[str, str]:
    return {"date": "2026-05-01", "quote_no": "Q-2026-001"}


def test_build_quote_creates_valid_docx(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    word.build_quote(
        out_path=out,
        vendor="AX상사",
        items=_sample_items(),
        meta=_sample_meta(),
    )
    assert out.exists()
    text = _read_full_text(out)
    assert "AX상사" in text
    assert "위젯" in text
    assert "견 적 서" in text or "견적서" in text


def test_build_quote_auto_total_correct(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    word.build_quote(
        out_path=out,
        vendor="V",
        items=_sample_items(),  # 10*50000 + 5*80000 = 900,000
        meta=_sample_meta(),
    )
    text = _read_full_text(out)
    assert "900,000" in text


def test_build_quote_explicit_total_overrides_auto(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    word.build_quote(
        out_path=out,
        vendor="V",
        items=_sample_items(),
        meta=_sample_meta(),
        total=1_000_000,
    )
    text = _read_full_text(out)
    assert "1,000,000" in text
    # 자동 합계 900,000은 표 셀에 등장할 수 있으나 "합 계" 줄에는 1,000,000이 들어가야 함.
    # 합계 단락 식별: "합 계:" 접두 단락에 1,000,000이 포함되었는지 확인.
    doc = Document(str(out))
    sum_paragraphs = [p.text for p in doc.paragraphs if "합 계" in p.text]
    assert sum_paragraphs, "expected '합 계' paragraph"
    assert any("1,000,000" in p for p in sum_paragraphs)


def test_build_quote_empty_items_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError, match="empty"):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=[],
            meta=_sample_meta(),
        )


def test_build_quote_meta_missing_date_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(KeyError, match="date"):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=_sample_items(),
            meta={"quote_no": "Q1"},
        )


def test_build_quote_meta_missing_quote_no_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(KeyError, match="quote_no"):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=_sample_items(),
            meta={"date": "2026-05-01"},
        )


def test_build_quote_column_map_renames(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    items: list[dict[str, object]] = [
        {"product_name": "위젯", "quantity": 1, "unit_price": 100},
    ]
    word.build_quote(
        out_path=out,
        vendor="V",
        items=items,
        meta=_sample_meta(),
        column_map={"name": "product_name", "qty": "quantity", "price": "unit_price"},
    )
    text = _read_full_text(out)
    assert "위젯" in text
    assert "100" in text


def test_build_quote_invalid_qty_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=[{"name": "x", "qty": "not_int", "price": 100}],
            meta=_sample_meta(),
        )


def test_build_quote_invalid_price_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=[{"name": "x", "qty": 1, "price": None}],
            meta=_sample_meta(),
        )


def test_build_quote_korean_text_preserved(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    items: list[dict[str, object]] = [
        {"name": "한글품목입니다", "qty": 2, "price": 12345},
    ]
    word.build_quote(
        out_path=out,
        vendor="에이엑스상사 주식회사",
        items=items,
        meta=_sample_meta(),
    )
    # 다시 읽어 한글 깨짐 없는지 확인
    text = _read_full_text(out)
    assert "한글품목입니다" in text
    assert "에이엑스상사 주식회사" in text


def test_build_quote_total_zero_when_all_zero_priced(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    items: list[dict[str, object]] = [
        {"name": "샘플", "qty": 5, "price": 0},
    ]
    word.build_quote(
        out_path=out,
        vendor="V",
        items=items,
        meta=_sample_meta(),
    )
    # 합계 0원이 표시되어야 함
    doc = Document(str(out))
    sum_paragraphs = [p.text for p in doc.paragraphs if "합 계" in p.text]
    assert sum_paragraphs
    assert any("0" in p for p in sum_paragraphs)


def test_build_quote_negative_qty_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError, match="negative"):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=[{"name": "x", "qty": -1, "price": 100}],
            meta=_sample_meta(),
        )


def test_build_quote_negative_price_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError, match="negative"):
        word.build_quote(
            out_path=out,
            vendor="V",
            items=[{"name": "x", "qty": 1, "price": -100}],
            meta=_sample_meta(),
        )


def test_build_quote_empty_vendor_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError, match="vendor"):
        word.build_quote(
            out_path=out,
            vendor="",
            items=_sample_items(),
            meta=_sample_meta(),
        )


def test_build_quote_whitespace_vendor_raises(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    with pytest.raises(ValueError, match="vendor"):
        word.build_quote(
            out_path=out,
            vendor="   ",
            items=_sample_items(),
            meta=_sample_meta(),
        )


def _collect_east_asia_fonts(path: Path) -> set[str]:
    """docx 내 모든 run의 w:rFonts/@w:eastAsia 값을 수집."""
    doc = Document(str(path))
    fonts: set[str] = set()
    # 단락
    for p in doc.paragraphs:
        for run in p.runs:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is None:
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                continue
            v = rFonts.get(qn("w:eastAsia"))
            if v:
                fonts.add(v)
    # 테이블 셀 단락
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        rPr = run._element.find(qn("w:rPr"))
                        if rPr is None:
                            continue
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            continue
                        v = rFonts.get(qn("w:eastAsia"))
                        if v:
                            fonts.add(v)
    return fonts


def test_build_quote_korean_font_applied(tmp_path: Path) -> None:
    out = tmp_path / "quote.docx"
    word.build_quote(
        out_path=out,
        vendor="AX상사",
        items=_sample_items(),
        meta=_sample_meta(),
    )
    fonts = _collect_east_asia_fonts(out)
    assert word.KOREAN_FONT in fonts, (
        f"expected eastAsia font {word.KOREAN_FONT!r} not found; got {fonts!r}"
    )


def test_build_quote_korean_font_env_override(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    custom = "Malgun Gothic"
    monkeypatch.setenv("AX_KOREAN_FONT", custom)
    # 모듈 상수는 import 시점 평가이므로, 런타임 override를 위해 모듈을 reload
    import importlib

    importlib.reload(word)
    try:
        out = tmp_path / "quote.docx"
        word.build_quote(
            out_path=out,
            vendor="AX상사",
            items=_sample_items(),
            meta=_sample_meta(),
        )
        fonts = _collect_east_asia_fonts(out)
        assert custom in fonts, f"expected {custom!r}; got {fonts!r}"
    finally:
        # 다른 테스트 격리 — 기본값으로 reload
        monkeypatch.delenv("AX_KOREAN_FONT", raising=False)
        importlib.reload(word)
