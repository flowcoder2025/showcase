"""Tests for case05 — 견적서/거래명세서 자동 생성 (docx + pdf) — T38 ScenarioResult."""

from pathlib import Path
from typing import Any

import pandas as pd
import pytest

_INPUT_NAME = "quote_requests.xlsx"


class _CaptureLogger:
    def __init__(self) -> None:
        self.infos: list[str] = []
        self.successes: list[str] = []
        self.warnings: list[str] = []
        self.errors: list[str] = []

    def info(self, msg: str) -> None:
        self.infos.append(msg)

    def success(self, msg: str) -> None:
        self.successes.append(msg)

    def warning(self, msg: str) -> None:
        self.warnings.append(msg)

    def error(self, msg: str) -> None:
        self.errors.append(msg)


def _stub_pdf(md_path: Path | str, out_path: Path | str, **_kw: Any) -> None:
    Path(out_path).write_bytes(b"%PDF-1.4\n%stub")


def _make_quote_df(n_requests: int = 10, items_per: int = 4) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for i in range(1, n_requests + 1):
        for j in range(items_per):
            rows.append(
                {
                    "견적번호": f"Q-2026-{i:03d}",
                    "거래처명": f"거래처{i:02d}",
                    "담당자": f"담당자{i}",
                    "이메일": f"v{i}@example.com",
                    "품목": f"부품-{i}-{j}",
                    "수량": 10 + j,
                    "단가": 50_000 + j * 1_000,
                    "납기일": "2026-06-30",
                }
            )
    return pd.DataFrame(rows)


def _write_input(in_dir: Path, df: pd.DataFrame) -> Path:
    in_dir.mkdir(parents=True, exist_ok=True)
    df.to_excel(in_dir / _INPUT_NAME, index=False)
    return in_dir


@pytest.fixture
def quote_input_dir(tmp_path: Path) -> Path:
    return _write_input(tmp_path / "in", _make_quote_df(n_requests=10, items_per=4))


def test_run_creates_docx_and_pdf_per_request(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)
    out = tmp_path / "out"

    result = scenario.run(input_dir=quote_input_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 10
    assert result["metrics"]["pdf_count"] == 10
    assert result["metrics"]["errors"] == 0
    assert len(list(out.glob("*.docx"))) == 10
    assert len(list(out.glob("*.pdf"))) == 10


def test_run_pdf_failure_does_not_block_other_requests(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    counter = {"n": 0}

    def flaky(md_path: Path | str, out_path: Path | str, **_kw: Any) -> None:
        counter["n"] += 1
        if counter["n"] in (3, 7):
            raise pdf_mod.MdToPdfError(f"simulated failure #{counter['n']}")
        Path(out_path).write_bytes(b"%PDF-1.4\n%stub")

    monkeypatch.setattr(pdf_mod, "md_to_pdf", flaky)
    out = tmp_path / "out"

    result = scenario.run(input_dir=quote_input_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 10
    assert result["metrics"]["pdf_count"] == 8
    assert result["metrics"]["errors"] >= 1


def test_run_uses_column_map_for_alternate_schema(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    rows: list[dict[str, Any]] = []
    for i in range(1, 4):
        for j in range(3):
            rows.append(
                {
                    "request_no": f"R-{i}",
                    "customer": f"Customer-{i}",
                    "product": f"prod-{j}",
                    "units": 5,
                    "unit_price": 100_000,
                    "delivery": "2026-07-01",
                }
            )
    in_dir = _write_input(tmp_path / "in", pd.DataFrame(rows))

    out = tmp_path / "out"
    result = scenario.run(
        input_dir=in_dir,
        output_dir=out,
        config={
            "column_map": {
                "request_id": "request_no",
                "vendor": "customer",
                "name": "product",
                "qty": "units",
                "price": "unit_price",
                "due_date": "delivery",
            }
        },
    )

    assert result["metrics"]["docx_count"] == 3
    assert result["metrics"]["pdf_count"] == 3
    assert result["metrics"]["errors"] == 0


def test_run_with_zero_requests(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    df = pd.DataFrame(
        columns=["견적번호", "거래처명", "담당자", "이메일", "품목", "수량", "단가", "납기일"]
    )
    in_dir = _write_input(tmp_path / "in", df)
    out = tmp_path / "out"

    result = scenario.run(input_dir=in_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 0
    assert result["metrics"]["pdf_count"] == 0
    assert result["metrics"]["errors"] == 0
    assert result["extras"]["requests"] == []


def test_run_summary_structure(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)
    out = tmp_path / "out"

    result = scenario.run(input_dir=quote_input_dir, output_dir=out)
    metrics = result["metrics"]
    requests = result["extras"]["requests"]

    for key in ("docx_count", "pdf_count", "errors"):
        assert key in metrics, f"missing key: {key}"
    assert isinstance(requests, list)
    assert len(requests) == 10
    for entry in requests:
        assert "request_id" in entry
        assert "vendor" in entry
        assert "n_items" in entry
        assert isinstance(entry["n_items"], int)
        assert entry["n_items"] >= 1


def test_run_creates_output_dir_if_missing(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    target = tmp_path / "nested" / "deeper" / "out"
    assert not target.exists()

    result = scenario.run(input_dir=quote_input_dir, output_dir=target)

    assert target.exists()
    assert target.is_dir()
    assert result["metrics"]["docx_count"] == 10


def test_run_docx_contains_vendor_and_items(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from docx import Document
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)
    out = tmp_path / "out"

    scenario.run(input_dir=quote_input_dir, output_dir=out)

    docx_files = sorted(out.glob("*.docx"))
    assert docx_files, "no docx generated"
    target = docx_files[0]
    doc = Document(str(target))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells
    )
    combined = full_text + "\n" + table_text

    assert "거래처01" in combined
    assert "견적번호" in combined
    assert "부품-1-" in combined


def test_run_continues_after_word_failure(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod
    from flowcoder_office_tools.docgen import word as word_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    real_build = word_mod.build_quote
    counter = {"n": 0}

    def flaky_build(**kwargs: Any) -> None:
        counter["n"] += 1
        if counter["n"] == 2:
            raise ValueError("simulated word failure")
        real_build(**kwargs)

    monkeypatch.setattr(word_mod, "build_quote", flaky_build)

    in_dir = _write_input(tmp_path / "in", _make_quote_df(n_requests=4, items_per=3))
    out = tmp_path / "out"

    result = scenario.run(input_dir=in_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 3
    assert result["metrics"]["errors"] >= 1
    assert result["metrics"]["pdf_count"] >= 3


def test_run_logs_per_request_progress(
    quote_input_dir: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    cap = _CaptureLogger()
    monkeypatch.setattr(scenario, "demo_logger", lambda _case: cap)

    out = tmp_path / "out"
    result = scenario.run(input_dir=quote_input_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 10
    request_lines = [m for m in cap.infos if "Q-2026-" in m]
    assert len(request_lines) == 10, f"expected 10 progress lines, got {len(request_lines)}"
    sample = request_lines[0]
    assert "Q-2026-001" in sample
    assert "거래처01" in sample
    assert "4" in sample


def test_run_progress_log_escapes_markup_in_vendor(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    rows: list[dict[str, Any]] = [
        {
            "견적번호": "Q-X-001",
            "거래처명": "[bold]회사[/bold]",
            "담당자": "담당",
            "이메일": "x@example.com",
            "품목": "부품-1",
            "수량": 1,
            "단가": 10_000,
            "납기일": "2026-06-30",
        }
    ]
    in_dir = _write_input(tmp_path / "in", pd.DataFrame(rows))

    cap = _CaptureLogger()
    monkeypatch.setattr(scenario, "demo_logger", lambda _case: cap)

    out = tmp_path / "out"
    scenario.run(input_dir=in_dir, output_dir=out)

    progress = [m for m in cap.infos if "Q-X-001" in m]
    assert progress, f"no progress log found, infos={cap.infos!r}"
    line = progress[0]
    assert "\\[bold]" in line, f"expected escaped bracket form in: {line!r}"
    assert "\\[/bold]" in line, f"expected escaped closing tag in: {line!r}"


def test_run_skips_empty_vendor_with_warning(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    rows: list[dict[str, Any]] = []
    rows.append(
        {
            "견적번호": "Q-EMPTY-001",
            "거래처명": "",
            "담당자": "x",
            "이메일": "x@example.com",
            "품목": "부품-1",
            "수량": 1,
            "단가": 10_000,
            "납기일": "2026-06-30",
        }
    )
    rows.append(
        {
            "견적번호": "Q-OK-002",
            "거래처명": "정상거래처",
            "담당자": "y",
            "이메일": "y@example.com",
            "품목": "부품-2",
            "수량": 2,
            "단가": 20_000,
            "납기일": "2026-06-30",
        }
    )
    in_dir = _write_input(tmp_path / "in", pd.DataFrame(rows))

    cap = _CaptureLogger()
    monkeypatch.setattr(scenario, "demo_logger", lambda _case: cap)

    out = tmp_path / "out"
    result = scenario.run(input_dir=in_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 1
    assert result["metrics"]["errors"] >= 1
    matched = [w for w in cap.warnings if "Q-EMPTY-001" in w]
    assert matched, f"expected warning mentioning Q-EMPTY-001, got {cap.warnings!r}"


def test_run_skips_whitespace_vendor(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    rows: list[dict[str, Any]] = [
        {
            "견적번호": "Q-WS-001",
            "거래처명": "   ",
            "담당자": "x",
            "이메일": "x@example.com",
            "품목": "부품-1",
            "수량": 1,
            "단가": 10_000,
            "납기일": "2026-06-30",
        },
        {
            "견적번호": "Q-OK-002",
            "거래처명": "정상거래처",
            "담당자": "y",
            "이메일": "y@example.com",
            "품목": "부품-2",
            "수량": 2,
            "단가": 20_000,
            "납기일": "2026-06-30",
        },
    ]
    in_dir = _write_input(tmp_path / "in", pd.DataFrame(rows))

    cap = _CaptureLogger()
    monkeypatch.setattr(scenario, "demo_logger", lambda _case: cap)

    out = tmp_path / "out"
    result = scenario.run(input_dir=in_dir, output_dir=out)

    assert result["metrics"]["docx_count"] == 1
    assert result["metrics"]["errors"] >= 1
    matched = [w for w in cap.warnings if "Q-WS-001" in w]
    assert matched, f"expected warning mentioning Q-WS-001, got {cap.warnings!r}"


def test_run_error_log_contains_request_id(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from flowcoder_office_tools.docgen import pdf as pdf_mod
    from flowcoder_office_tools.docgen import word as word_mod

    from cases.case05_doc_quote_generator import scenario

    monkeypatch.setattr(pdf_mod, "md_to_pdf", _stub_pdf)

    real_build = word_mod.build_quote
    counter = {"n": 0}

    def flaky_build(**kwargs: Any) -> None:
        counter["n"] += 1
        if counter["n"] == 2:
            raise RuntimeError("boom")
        real_build(**kwargs)

    monkeypatch.setattr(word_mod, "build_quote", flaky_build)

    cap = _CaptureLogger()
    monkeypatch.setattr(scenario, "demo_logger", lambda _case: cap)

    in_dir = _write_input(tmp_path / "in", _make_quote_df(n_requests=3, items_per=2))
    out = tmp_path / "out"

    result = scenario.run(input_dir=in_dir, output_dir=out)

    assert result["metrics"]["errors"] >= 1
    matched = [w for w in cap.warnings if "Q-2026-002" in w and "RuntimeError" in w]
    assert matched, f"expected warning with request_id and RuntimeError, got {cap.warnings!r}"
