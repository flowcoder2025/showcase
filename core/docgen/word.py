"""견적서/거래명세서 생성 — python-docx.

case05 docgen 카테고리의 핵심 빌더. column_map 강제로 다른 입력 스키마와도
호환되며, total 미명시 시 sum(qty*price)을 자동 계산한다.

NOTE: 한글 폰트는 시스템 기본(macOS: Apple SD Gothic Neo)에 의존한다.
명시적 폰트 설정은 T4.5 후속 fixer 또는 case05 wrapper에서 보강 가능.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

DEFAULT_COLUMN_MAP: dict[str, str] = {"name": "name", "qty": "qty", "price": "price"}


def _coerce_int(value: Any, *, field: str, item: dict[str, Any]) -> int:
    """qty/price 값을 int로 강제 변환 — 실패 시 ValueError."""
    if value is None:
        raise ValueError(f"item {item!r} has invalid {field}: value is None")
    try:
        return int(value)
    except (TypeError, ValueError) as e:
        raise ValueError(f"item {item!r} has invalid {field}: {e}") from e


def build_quote(
    *,
    out_path: Path | str,
    vendor: str,
    items: list[dict[str, Any]],
    meta: dict[str, str],
    column_map: dict[str, str] | None = None,
    total: int | None = None,
) -> None:
    """견적서 docx 생성.

    Args:
        out_path: 결과 .docx 경로
        vendor: 거래처명
        items: 품목 리스트 — 각 dict는 column_map 기준 name/qty/price 키 보유
        meta: {"date": YYYY-MM-DD, "quote_no": "Q-XXX-NNN", ...} (date/quote_no 필수)
        column_map: 입력 dict 키 매핑 (재사용성 — 다른 입력 스키마 호환)
        total: 명시 시 사용, 미명시 시 sum(qty * price) 자동 계산

    Raises:
        ValueError: items 빈 리스트 또는 qty/price 변환 실패
        KeyError: meta에 'date' 또는 'quote_no' 누락, items에 매핑된 키 누락
    """
    if not items:
        raise ValueError("items must not be empty")
    if "date" not in meta:
        raise KeyError("meta must contain 'date'")
    if "quote_no" not in meta:
        raise KeyError("meta must contain 'quote_no'")

    cm = {**DEFAULT_COLUMN_MAP, **(column_map or {})}
    name_key = cm["name"]
    qty_key = cm["qty"]
    price_key = cm["price"]

    # 자동 합계 (total 미명시 시) — 동시에 모든 row의 qty/price 검증
    line_amounts: list[int] = []
    for it in items:
        qty = _coerce_int(it.get(qty_key), field="qty", item=it)
        price = _coerce_int(it.get(price_key), field="price", item=it)
        line_amounts.append(qty * price)
    auto_total = sum(line_amounts)
    final_total = total if total is not None else auto_total

    doc = Document()
    # 헤더
    doc.add_heading("견 적 서", level=0)
    # 메타 단락
    doc.add_paragraph(f"견적번호: {meta['quote_no']}")
    doc.add_paragraph(f"작성일: {meta['date']}")
    doc.add_paragraph(f"거래처: {vendor}")

    # 품목 표
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "품목"
    hdr[1].text = "수량"
    hdr[2].text = "단가"
    hdr[3].text = "금액"
    for i, it in enumerate(items, start=1):
        # 위에서 이미 검증했으므로 여기서는 안전하게 int() 재변환
        qty = int(it[qty_key])
        price = int(it[price_key])
        row = table.rows[i].cells
        row[0].text = str(it[name_key])
        row[1].text = f"{qty:,}"
        row[2].text = f"{price:,}"
        row[3].text = f"{qty * price:,}"

    # 합계 — 굵게
    doc.add_paragraph()
    total_para = doc.add_paragraph(f"합 계: {final_total:,}원")
    total_para.runs[0].bold = True

    doc.save(str(out_path))
